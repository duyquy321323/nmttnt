"""Extract text from txt / docx / pdf"""

import io
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.core.constants import ALLOWED_DOCUMENT_EXTENSIONS


class DocumentParser:
    """アップロードファイルからプレーンテキストを抽出"""

    @staticmethod
    def validate_extension(filename: str) -> str:
        suffix = Path(filename).suffix.lower()
        if suffix not in ALLOWED_DOCUMENT_EXTENSIONS:
            allowed = ", ".join(sorted(ALLOWED_DOCUMENT_EXTENSIONS))
            raise ValueError(f"Định dạng không hỗ trợ. Chỉ chấp nhận: {allowed}")
        return suffix

    def parse(self, filename: str, content: bytes) -> str:
        suffix = self.validate_extension(filename)
        if suffix == ".txt":
            return self._parse_txt(content)
        if suffix == ".docx":
            return self._parse_docx(content)
        return self._parse_pdf(content)

    @staticmethod
    def _parse_txt(content: bytes) -> str:
        for encoding in ("utf-8", "utf-8-sig", "cp1258", "latin-1"):
            try:
                return content.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        raise ValueError("Không đọc được file txt (mã hóa không hỗ trợ).")

    @staticmethod
    def _parse_docx(content: bytes) -> str:
        """段落と表セルからテキストを抽出（教材docxは表に本文があることが多い）"""
        document = Document(io.BytesIO(content))
        parts: list[str] = []

        for paragraph in document.paragraphs:
            line = paragraph.text.strip()
            if line:
                parts.append(line)

        for table in document.tables:
            for row in table.rows:
                row_lines: list[str] = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and (not row_lines or cell_text != row_lines[-1]):
                        row_lines.append(cell_text)
                if row_lines:
                    parts.append("\n".join(row_lines))

        text = "\n".join(parts).strip()
        if not text:
            raise ValueError("File docx không có nội dung văn bản.")
        return text

    @staticmethod
    def _parse_pdf(content: bytes) -> str:
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages.append(page_text.strip())
        text = "\n".join(pages).strip()
        if not text:
            raise ValueError("File pdf không trích xuất được văn bản.")
        return text
